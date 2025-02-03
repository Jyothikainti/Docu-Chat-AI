import re
from io import BytesIO
from typing import List, Tuple
from langchain.docstore.document import Document
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores.faiss import FAISS
from pypdf import PdfReader
from docx import Document as DocxDocument

# Parse PDF and extract text
def parse_pdf(file: BytesIO, filename: str) -> Tuple[List[str], str]:
    pdf = PdfReader(file)
    text_pages = []
    for page in pdf.pages:
        text = page.extract_text()
        text = re.sub(r"(\w+)-\n(\w+)", r"\1\2", text)
        text = re.sub(r"(?<!\n\s)\n(?!\s\n)", " ", text.strip())
        text_pages.append(text)
    return text_pages, filename

# Parse DOCX and extract text
def parse_docx(file: BytesIO, filename: str) -> Tuple[List[str], str]:
    doc = DocxDocument(file)
    text_pages = [para.text for para in doc.paragraphs if para.text.strip()]
    return text_pages, filename

# Convert text to LangChain documents
def text_to_docs(text: List[str], filename: str) -> List[Document]:
    if isinstance(text, str):
        text = [text]
    page_docs = [Document(page_content=page) for page in text]
    for i, doc in enumerate(page_docs):
        doc.metadata["page"] = i + 1

    chunks = []
    splitter = RecursiveCharacterTextSplitter(chunk_size=4000, chunk_overlap=0)
    for doc in page_docs:
        split_chunks = splitter.split_text(doc.page_content)
        for i, chunk in enumerate(split_chunks):
            chunk_doc = Document(
                page_content=chunk,
                metadata={"page": doc.metadata["page"], "chunk": i, "filename": filename},
            )
            chunks.append(chunk_doc)
    return chunks

# Convert documents to FAISS index
def docs_to_index(docs, api_key):
    return FAISS.from_documents(docs, OpenAIEmbeddings(openai_api_key=api_key))

# Generate index for PDFs and DOCX files
def get_index_for_files(files, filenames, api_key):
    docs = []
    for file, filename in zip(files, filenames):
        if filename.endswith(".pdf"):
            text, filename = parse_pdf(BytesIO(file), filename)
        elif filename.endswith(".docx"):
            text, filename = parse_docx(BytesIO(file), filename)
        else:
            continue
        docs.extend(text_to_docs(text, filename))
    return docs_to_index(docs, api_key)
